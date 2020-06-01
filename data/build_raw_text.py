import docx
import os
from functional import seq
import re

docx_dir = r'D:\zqcyks'
docxes = []

for file in os.listdir(docx_dir):
    docxes.append(docx.Document(os.path.join(docx_dir, file)))


def remove_blank(s):
    return re.sub('\s+', '', s).strip()


def read_paragraph(x):
    lines = []
    for para in x:
        lines.append(remove_blank(para.text))
    return lines


def read_table(x):
    lines = []
    for tb in x:
        row_count = len(tb.rows)
        col_count = len(tb.columns)

        for i in range(row_count):
            for j in range(col_count):
                lines.append(remove_blank(tb.cell(i, j).text))

    return lines


def read_docx(x):
    return read_paragraph(x.paragraphs) + read_table(x.tables)


raw_text = (seq(docxes)
            .map(read_docx)
            .flatten()
            .filter(lambda x: len(x) > 10)
            ).list()

with open('raw_text.txt', 'w', encoding='utf-8') as f:
    for line in raw_text:
        f.write(line + '\n')
